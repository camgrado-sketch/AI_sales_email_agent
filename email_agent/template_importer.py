import hashlib
import html
import json
import os
import re
import shutil
from datetime import datetime
from typing import Optional

from email_agent import config, data_store, llm_client, template_engine


class ImportCandidate:
    def __init__(self, path, checksum):
        self.path = path
        self.checksum = checksum
        self.filename = os.path.basename(path)
        self.ext = os.path.splitext(path)[1].lower()


def _ensure_dirs():
    os.makedirs(config.TEMPLATE_IMPORT_DIR, exist_ok=True)
    os.makedirs(config.TEMPLATE_ARCHIVE_DIR, exist_ok=True)


def _file_checksum(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _normalize_template_name(name):
    """Normalize a template name from a filename or user input."""
    name = re.sub(r"[^\w\-]", "_", name).strip("_").lower()
    return name or "initial_contact"


def _image_ext_from_content_type(content_type):
    """Map a MIME image type to a file extension."""
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/webp": ".webp",
        "image/tiff": ".tiff",
    }
    return mapping.get(content_type, ".png")


def _images_used_by_active_drafts(template_name=None):
    """Collect image filenames referenced by 待审核/待发送 drafts.

    待审核 = review_status == "pending"；待发送 = review_status == "approved"
    且尚无成功发送记录。这些草稿引用的图片资源在重新导入模板时不得删除，
    否则未发送的邮件将丢失图片。
    """
    protected = set()
    drafts = data_store.load_drafts()
    if not drafts:
        return protected
    sent_ids = data_store.get_sent_draft_ids()
    for draft in drafts:
        draft_template = draft.get("template")
        if (
            template_name
            and draft_template
            and draft_template != template_name
        ):
            continue
        review_status = draft.get("review_status")
        pending_review = review_status == "pending"
        pending_send = review_status == "approved" and (
            draft.get("draft_id") not in sent_ids
        )
        if not (pending_review or pending_send):
            continue
        for img in draft.get("images") or []:
            img_path = img.get("path", "") if isinstance(img, dict) else ""
            if img_path:
                protected.add(os.path.basename(img_path))
    return protected


def _cleanup_template_images(template_name):
    """Remove images from assets/images/ that belong to a previous import of template_name.

    被"待审核"或"待发送"草稿引用的图片禁止删除，避免未发送邮件丢失图片资源。
    """
    if not os.path.exists(config.IMAGES_DIR):
        return
    protected = _images_used_by_active_drafts(template_name)
    prefix = f"{template_name}_img_"
    kept = 0
    for fname in os.listdir(config.IMAGES_DIR):
        if not fname.startswith(prefix):
            continue
        if fname in protected:
            kept += 1
            continue
        os.remove(os.path.join(config.IMAGES_DIR, fname))
    if kept:
        print(
            f"⚠️ 有 {kept} 张图片正被待审核/待发送的草稿引用，已保留，"
            "避免邮件图片丢失。"
        )


def _template_name_from_filename(filename):
    """Map a source filename to a standard template name.

    Chinese filenames are mapped by keyword; any unmatched file falls back
    to the 'other' category so the original filename never appears as a
    terminal variable or directory name.
    """
    base = os.path.splitext(filename)[0].lower()
    if any(k in base for k in ("开发信", "首封", "initial", "first")):
        return "initial_contact"
    if any(k in base for k in ("跟进", "follow", "reminder", "second")):
        return "follow_up"
    if any(k in base for k in ("最终", "final", "last", "收尾")):
        return "final_note"
    return "other"


# ------------------------------------------------------------------------------
# Scanning and change detection
# ------------------------------------------------------------------------------

def scan_import_folder():
    """Return a list of ImportCandidate for supported files in the import folder."""
    _ensure_dirs()
    candidates = []
    for root, _dirs, files in os.walk(config.TEMPLATE_IMPORT_DIR):
        for fname in files:
            ext = os.path.splitext(fname)[1].lower()
            if ext not in (".md", ".docx", ".pdf"):
                continue
            path = os.path.join(root, fname)
            candidates.append(ImportCandidate(path, _file_checksum(path)))
    return candidates


def detect_changes():
    """Return import candidates whose checksum differs from the stored state."""
    state = data_store.load_template_import_state()
    stored = state.get("files", {})
    changed = []
    for cand in scan_import_folder():
        rel = os.path.relpath(cand.path, config.TEMPLATE_IMPORT_DIR)
        if stored.get(rel) != cand.checksum:
            changed.append(cand)
    return changed


def save_import_state():
    """Snapshot current import folder checksums."""
    state = data_store.load_template_import_state()
    state["last_scan_at"] = datetime.now().isoformat()
    state["files"] = {
        os.path.relpath(cand.path, config.TEMPLATE_IMPORT_DIR): cand.checksum
        for cand in scan_import_folder()
    }
    data_store.save_template_import_state(state)


# ------------------------------------------------------------------------------
# Content extraction -> Markdown
# ------------------------------------------------------------------------------

def extract_to_markdown(path, template_name=None):
    """Extract a Markdown-like text representation from md/docx/pdf."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".md":
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    if ext == ".docx":
        return _docx_to_markdown(path, template_name=template_name)
    if ext == ".pdf":
        return _pdf_to_markdown(path)
    raise ValueError(f"Unsupported template file format: {ext}")


def _docx_to_markdown(path, template_name=None):
    try:
        import docx
        from docx.oxml.ns import qn
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for .docx import. Run: pip install python-docx"
        ) from e

    doc = docx.Document(path)
    os.makedirs(config.IMAGES_DIR, exist_ok=True)

    base_name = template_name or _template_name_from_filename(os.path.basename(path)) or "template"
    image_counter = 0
    extracted_partnames = set()

    def resolve_image_part(rel_id):
        """Resolve an image part by relationship id.

        Primary lookup is doc.part.related_parts; the rels-table fallback
        covers embedding styles produced by different Word versions where
        the part is not surfaced through related_parts.
        """
        if not rel_id:
            return None
        image_part = doc.part.related_parts.get(rel_id)
        if image_part is not None:
            return image_part
        rel = doc.part.rels.get(rel_id)
        if rel is not None and not getattr(rel, "is_external", True):
            return rel.target_part
        return None

    def save_image_part(image_part):
        nonlocal image_counter
        image_counter += 1
        image_name = f"{base_name}_img_{image_counter:02d}"
        ext = _image_ext_from_content_type(image_part.content_type)
        image_path = os.path.join(config.IMAGES_DIR, f"{image_name}{ext}")
        with open(image_path, "wb") as f:
            f.write(image_part.blob)
        extracted_partnames.add(str(image_part.partname))
        return image_name

    drawing_tag = qn("w:drawing")
    drawingml_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    blip_xpath = f".//{{{drawingml_ns}}}blip"
    imagedata_xpath = ".//{urn:schemas-microsoft-com:vml}imagedata"

    lines = []
    for para in doc.paragraphs:
        style = (para.style.name or "").lower()
        parts = []
        for run in para.runs:
            run_text = run.text
            # Descendant search: newer Word versions wrap drawings in
            # mc:AlternateContent, so direct-child lookup would miss them.
            drawings = run._element.findall(".//" + drawing_tag)
            # Legacy VML images (w:pict / v:imagedata) only matter when the run
            # has no modern drawing; otherwise the mc:Fallback copy would be a
            # duplicate of the same image.
            imagedata_list = [] if drawings else run._element.findall(imagedata_xpath)
            if not drawings and not imagedata_list:
                if run_text:
                    parts.append(run_text)
                continue

            # Preserve any text in the same run as the image.
            if run_text:
                parts.append(run_text)

            if drawings:
                for drawing in drawings:
                    # Every blip (group shapes contain several).
                    for blip in drawing.findall(blip_xpath):
                        rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
                        image_part = resolve_image_part(rel_id)
                        if image_part is None:
                            continue
                        image_name = save_image_part(image_part)
                        parts.append(f"{{{{IMAGE:{image_name}}}}}")
            else:
                for imagedata in imagedata_list:
                    image_part = resolve_image_part(imagedata.get(qn("r:id")))
                    if image_part is None:
                        continue
                    image_name = save_image_part(image_part)
                    parts.append(f"{{{{IMAGE:{image_name}}}}}")

        para_text = "".join(parts).strip()
        if not para_text:
            continue
        if "heading" in style:
            level = re.search(r"\d", style)
            prefix = "#" * (int(level.group(0)) if level else 1)
            lines.append(f"{prefix} {para_text}")
        else:
            lines.append(para_text)

    # Safety net: walk every image part in the package so images anchored in
    # tables, text boxes, headers/footers, or other structures outside
    # doc.paragraphs are never silently lost. /docProps/ thumbnails are not
    # body content and are excluded.
    orphans = []
    for part in doc.part.package.iter_parts():
        content_type = getattr(part, "content_type", None) or ""
        if not content_type.startswith("image/"):
            continue
        partname = str(part.partname)
        if partname.startswith("/docProps/"):
            continue
        if partname in extracted_partnames:
            continue
        image_name = save_image_part(part)
        orphans.append(image_name)
    if orphans:
        print(
            f"⚠️ 警告：发现 {len(orphans)} 张未在正文段落中定位到的图片"
            f"（可能位于表格、文本框或页眉页脚），已追加到模板末尾："
            + "、".join(orphans)
        )
        lines.extend(f"{{{{IMAGE:{name}}}}}" for name in orphans)

    return "\n\n".join(lines)


def _pdf_to_markdown(path):
    try:
        import pdfplumber
    except ImportError as e:
        raise RuntimeError(
            "pdfplumber is required for PDF import. Run: pip install pdfplumber"
        ) from e

    paragraphs = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue
            # Split into paragraphs by blank lines
            for para in re.split(r"\n\s*\n", text.strip()):
                para = " ".join(line.strip() for line in para.splitlines() if line.strip())
                if para:
                    paragraphs.append(para)
    return "\n\n".join(paragraphs)


def _truncate_subject(text, max_len=60):
    """Truncate a subject line to max_len characters, preferably at a word boundary."""
    text = (text or "").strip()
    if len(text) <= max_len:
        return text
    truncated = text[:max_len]
    if " " in truncated:
        truncated = truncated.rsplit(" ", 1)[0]
    return truncated.rstrip()


def _fallback_subject(markdown):
    """Derive a non-empty subject line from markdown when the LLM returns none."""
    if not markdown:
        return "GRADO Contract Partnership Opportunity"

    # Prefer the first Markdown heading.
    for line in markdown.splitlines():
        line = line.strip()
        m = re.match(r"^#+\s+(.+)", line)
        if m:
            return _truncate_subject(m.group(1))

    # Otherwise use the first non-empty line.
    for line in markdown.splitlines():
        line = line.strip()
        if line:
            return _truncate_subject(line)

    return "GRADO Contract Partnership Opportunity"


# ------------------------------------------------------------------------------
# Language detection
# ------------------------------------------------------------------------------

def detect_language(markdown, filename=""):
    """Detect whether the markdown is primarily Chinese (cn) or English (en)."""
    base = os.path.splitext(filename)[0].lower()
    if any(k in base for k in ("cn", "chinese", "zh", "中文")):
        return "cn"
    if any(k in base for k in ("en", "english", "eng", "英文")):
        return "en"

    sample = markdown[:2000]
    try:
        from langdetect import detect
        lang = detect(sample)
        if lang.startswith("zh"):
            return "cn"
        if lang == "en":
            return "en"
    except Exception:
        pass

    # Fallback: CJK ratio
    cjk = len(re.findall(r"[一-鿿]", sample))
    total = len(sample.strip())
    if total == 0:
        return "en"
    return "cn" if cjk / total > 0.1 else "en"


# ------------------------------------------------------------------------------
# LLM-based template structuring
# ------------------------------------------------------------------------------

def _available_variable_names():
    """Return the canonical list of variable placeholders the LLM may use."""
    return [
        "SENDER_NAME",
        "SENDER_TITLE",
        "SENDER_COMPANY",
        "SENDER_EMAIL",
        "SENDER_PHONE",
        "SENDER_MARKET_REGION",
        "CUSTOMER_FIRST_NAME",
        "CUSTOMER_NAME",
        "CUSTOMER_COMPANY",
        "CUSTOMER_POSITION",
        "CUSTOMER_LOCATION",
        "CUSTOMER_INDUSTRY",
        "CURRENT_DATE",
    ]


def _template_structure_schema():
    return {
        "name": "structured_email_template",
        "type": "object",
        "properties": {
            "subject_template": {
                "type": "string",
                "description": "Email subject line with {{VAR}} placeholders",
            },
            "subject_template_cn": {
                "type": "string",
                "description": (
                    "Chinese subject line with {{VAR}} placeholders"
                ),
            },
            "subject_template_en": {
                "type": "string",
                "description": (
                    "English subject line with {{VAR}} placeholders; "
                    "MUST NOT contain any Chinese characters"
                ),
            },
            "cn_html": {
                "type": "string",
                "description": "Complete Chinese HTML email body with {{VAR}}, {{IMAGE:name}}, {{FILE:name}} placeholders",
            },
            "en_html": {
                "type": "string",
                "description": "Complete English HTML email body with placeholders",
            },
            "variables": {
                "type": "array",
                "items": {"type": "string"},
                "description": "List of variable names actually used in the templates",
            },
            "images": {
                "type": "array",
                "items": {"type": "string"},
                "description": "List of image placeholder names",
            },
            "files": {
                "type": "array",
                "items": {"type": "string"},
                "description": "List of file placeholder names",
            },
            "ignored_sections": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Sections that were ignored as non-core content",
            },
        },
        "required": ["subject_template", "cn_html", "en_html", "variables", "images", "files"],
        "additionalProperties": False,
    }


def structure_template_with_llm(markdown, filename=""):
    """Call the remote LLM once to turn extracted Markdown into a structured bilingual template.

    The LLM receives only text and placeholder names; no image or file binaries are uploaded.
    """
    if not config.get_active_model():
        raise RuntimeError("No active LLM model configured. Check .env or settings.")

    system_prompt = _read_file(config.TEMPLATE_IMPORT_PROMPT_FILE)
    if not system_prompt:
        raise RuntimeError(f"Template import prompt not found: {config.TEMPLATE_IMPORT_PROMPT_FILE}")

    user_prompt = f"""Source filename: {filename}

Available variables (use only these exact names):
{', '.join(_available_variable_names())}

---

Extracted content:

{markdown}
"""

    raw = llm_client.complete_json(
        system_prompt,
        user_prompt,
        _template_structure_schema(),
        temperature=0.3,
    )
    try:
        content = raw["content"]
    except (TypeError, KeyError) as e:
        raise RuntimeError(
            f"LLM 返回结果缺少 content 字段（{e}）。请重试导入，或在设置（菜单 S）中切换模型。"
        ) from e
    try:
        result = json.loads(content)
    except (TypeError, json.JSONDecodeError) as e:
        raise RuntimeError(
            f"LLM 返回内容无法解析为 JSON（{e}）。请重试导入，或在设置（菜单 S）中切换模型。"
        ) from e
    if not isinstance(result, dict):
        raise RuntimeError("LLM 返回的模板结构不是 JSON 对象，请重试导入或切换模型。")

    # Normalize arrays to strings and strip whitespace
    for key in ("subject_template", "cn_html", "en_html"):
        if key in result:
            result[key] = str(result[key]).strip()
    for key in ("variables", "images", "files", "ignored_sections"):
        value = result.get(key)
        if not isinstance(value, list):
            result[key] = []
        else:
            result[key] = [str(v).strip() for v in value if str(v).strip()]

    # 容错：LLM 可能返回命名不符的语言字段（如 chinese_html），归一化到 cn_html/en_html。
    for canonical, aliases in (
        ("cn_html", ("chinese_html", "zh_html", "html_cn")),
        ("en_html", ("english_html", "html_en")),
    ):
        if not str(result.get(canonical, "") or "").strip():
            for alt in aliases:
                alt_value = result.get(alt)
                if str(alt_value or "").strip():
                    result[canonical] = str(alt_value)
                    print(f"ℹ️ LLM 输出使用了字段名 '{alt}'，已归一化为 '{canonical}'。")
                    break

    # 容错日志：任一语言 HTML 缺失时给出明确中文提示，而不是让下游抛出难懂的错误。
    for key, label in (("cn_html", "中文"), ("en_html", "英文")):
        if not str(result.get(key, "") or "").strip():
            print(
                f"⚠️ 警告：LLM 输出缺少{label}模板内容（'{key}' 为空），"
                f"{label}版模板将不会生成。可稍后重新导入，或手工补充对应模板文件。"
            )

    # Guard: English HTML must not contain Chinese characters.
    en_html = result.get("en_html", "")
    if re.search(r"[一-鿿]", en_html):
        print("⚠️ 警告：en_html 中检测到汉字，请检查英文模板语言纯净度。")

    return result


# ------------------------------------------------------------------------------
# Write structured template to disk
# ------------------------------------------------------------------------------

def _default_config_yaml(template_name, structured, source_lang=None):
    """Build a default config.yaml content from the structured template."""
    variables = structured.get("variables", [])
    images = structured.get("images", [])
    files = structured.get("files", [])
    subject_template = structured.get("subject_template", "")

    lines = [
        f"template_name: {template_name}",
    ]
    if source_lang:
        lines.append(f"source_language: {source_lang}")
    lines.extend([
        "purpose: Auto-imported template",
        "customer_type: all",
        "recommended_stage: new_lead",
        f"subject_template: {subject_template!r}",
    ])
    if structured.get("subject_template_cn"):
        lines.append(
            f"subject_template_cn: {structured['subject_template_cn']!r}"
        )
    if structured.get("subject_template_en"):
        lines.append(
            f"subject_template_en: {structured['subject_template_en']!r}"
        )
    lines.append("variables:")
    for v in variables:
        lines.append(f"  - {v}")
    lines.append("images:")
    for img in images:
        lines.append(f"  - {img}")
    lines.append("files:")
    for f in files:
        lines.append(f"  - {f}")
    lines.extend([
        "rules:",
        "  - Use only facts from the customer record and confirmed sender profile.",
        "  - Do not invent names, companies, positions, or locations.",
        "  - The email language is determined by the language parameter at render time.",
    ])
    return "\n".join(lines) + "\n"


def write_structured_template(template_name, structured, source_lang):
    """Write the bilingual template files and config.yaml to the active template directory.

    Args:
        template_name: target template directory name.
        structured: dict returned by structure_template_with_llm().
        source_lang: "cn" or "en" — determines which version becomes template.html.

    Returns:
        dict with paths: main_path, other_path, config_path, source_language, target_language.
        other_path is None when the LLM did not produce the target-language HTML.

    Raises:
        ValueError: when the source-language HTML is missing, since an empty
            template.html would break the whole template.
    """
    target_lang = "en" if source_lang == "cn" else "cn"
    source_html_key = "cn_html" if source_lang == "cn" else "en_html"
    target_html_key = "en_html" if source_lang == "cn" else "cn_html"

    source_html = str(structured.get(source_html_key, "") or "")
    target_html = str(structured.get(target_html_key, "") or "")
    if not source_html.strip():
        raise ValueError(
            f"模板导入失败：LLM 未返回源语言（{source_lang}）模板内容"
            f"（'{source_html_key}' 为空）。请重试导入，或检查 LLM 配置。"
        )

    target_dir = os.path.join(config.TEMPLATES_DIR, template_name)
    os.makedirs(target_dir, exist_ok=True)

    main_path = os.path.join(target_dir, "template.html")
    other_path = os.path.join(target_dir, f"template_{target_lang}.html")
    config_path = os.path.join(target_dir, "config.yaml")

    with open(main_path, "w", encoding="utf-8") as f:
        f.write(source_html)

    written_other_path = None
    if target_html.strip():
        with open(other_path, "w", encoding="utf-8") as f:
            f.write(target_html)
        written_other_path = other_path
    else:
        print(
            f"⚠️ 警告：{target_lang} 版模板未写入（LLM 输出缺少 '{target_html_key}' 字段）。"
            f"可手工创建 {other_path} 或重新导入。"
        )

    with open(config_path, "w", encoding="utf-8") as f:
        f.write(_default_config_yaml(template_name, structured, source_lang))

    return {
        "main_path": main_path,
        "other_path": written_other_path,
        "config_path": config_path,
        "source_language": source_lang,
        "target_language": target_lang,
    }


# ------------------------------------------------------------------------------
# Archiving and activation
# ------------------------------------------------------------------------------

def archive_current_template(template_name):
    """Copy the current active template directory to templates/archive/<template_name>/YYYY/MM/DD/."""
    src = os.path.join(config.TEMPLATES_DIR, template_name)
    if not os.path.exists(src):
        return None
    now = datetime.now()
    stamp = now.strftime("%H%M%S")
    dst = os.path.join(
        config.TEMPLATE_ARCHIVE_DIR,
        template_name,
        now.strftime("%Y"),
        now.strftime("%m"),
        now.strftime("%d"),
        stamp,
    )
    shutil.copytree(src, dst)
    return dst


def list_archive_folders():
    """Return a flat list of archived template folders sorted newest first.

    Supports both the current layout:
        archive/<template_name>/YYYY/MM/DD/<stamp>/
    and the legacy layout:
        archive/YYYY/MM/DD/<template_name>_<stamp>/
    """
    archives = []
    base = config.TEMPLATE_ARCHIVE_DIR
    if not os.path.exists(base):
        return archives

    # New structure: archive/<template_name>/YYYY/MM/DD/<stamp>
    for template_name in sorted(os.listdir(base), reverse=True):
        tpath = os.path.join(base, template_name)
        if not os.path.isdir(tpath) or template_name in (".gitkeep",):
            continue
        if re.match(r"^\d{4}$", template_name):
            continue  # legacy top-level year
        for year in sorted(os.listdir(tpath), reverse=True):
            ypath = os.path.join(tpath, year)
            if not os.path.isdir(ypath) or not re.match(r"^\d{4}$", year):
                continue
            for month in sorted(os.listdir(ypath), reverse=True):
                mpath = os.path.join(ypath, month)
                if not os.path.isdir(mpath) or not re.match(r"^\d{2}$", month):
                    continue
                for day in sorted(os.listdir(mpath), reverse=True):
                    dpath = os.path.join(mpath, day)
                    if not os.path.isdir(dpath) or not re.match(r"^\d{2}$", day):
                        continue
                    for stamp in sorted(os.listdir(dpath), reverse=True):
                        epath = os.path.join(dpath, stamp)
                        if os.path.isdir(epath) and re.match(r"^\d{6}$", stamp):
                            archives.append({
                                "path": epath,
                                "date": f"{year}/{month}/{day}",
                                "name": f"{template_name}/{year}/{month}/{day}/{stamp}",
                                "template_name": template_name,
                                "stamp": stamp,
                            })

    # Old structure: archive/YYYY/MM/DD/<name>_<stamp>
    for year in sorted(os.listdir(base), reverse=True):
        ypath = os.path.join(base, year)
        if not os.path.isdir(ypath) or not re.match(r"^\d{4}$", year):
            continue
        for month in sorted(os.listdir(ypath), reverse=True):
            mpath = os.path.join(ypath, month)
            if not os.path.isdir(mpath) or not re.match(r"^\d{2}$", month):
                continue
            for day in sorted(os.listdir(mpath), reverse=True):
                dpath = os.path.join(mpath, day)
                if not os.path.isdir(dpath) or not re.match(r"^\d{2}$", day):
                    continue
                for name in sorted(os.listdir(dpath), reverse=True):
                    epath = os.path.join(dpath, name)
                    if not os.path.isdir(epath):
                        continue
                    template_name, stamp = name, ""
                    if "_" in name:
                        parts = name.rsplit("_", 1)
                        if parts[1].isdigit() and len(parts[1]) == 6:
                            template_name, stamp = parts[0], parts[1]
                    archives.append({
                        "path": epath,
                        "date": f"{year}/{month}/{day}",
                        "name": name,
                        "template_name": template_name,
                        "stamp": stamp,
                    })

    archives.sort(key=lambda x: (x["date"].replace("/", ""), x["stamp"]), reverse=True)
    return archives


def get_active_template_path(template_name):
    """Return the active template directory path if it exists."""
    path = os.path.join(config.TEMPLATES_DIR, template_name)
    return path if os.path.exists(path) else None


def get_latest_archive(template_name):
    """Return the newest archive path for a template, or None."""
    matches = [a for a in list_archive_folders() if a["template_name"] == template_name]
    return matches[0]["path"] if matches else None


def list_archives_by_day(template_name):
    """Return [(date, [archive_dict, ...]), ...] sorted by date descending."""
    matches = [a for a in list_archive_folders() if a["template_name"] == template_name]
    days = {}
    for a in matches:
        days.setdefault(a["date"], []).append(a)
    return sorted(days.items(), key=lambda x: x[0].replace("/", ""), reverse=True)


def list_template_names_in_archive():
    """Return sorted list of template names that have archives."""
    return sorted({a["template_name"] for a in list_archive_folders()})


def is_import_state_stale():
    """Return True when templates/email/ is empty but templates/import/ still
    contains files already recorded in template_import_state.json."""
    if template_engine.list_templates():
        return False
    candidates = scan_import_folder()
    if not candidates:
        return False
    state = data_store.load_template_import_state()
    return bool(state.get("files"))


def reset_import_state():
    """Clear the checksum state so existing import files appear as new."""
    data_store.save_template_import_state({
        "last_reset_at": datetime.now().isoformat(),
        "files": {},
    })


def delete_archive_entry(path):
    """Remove an archived template folder tree."""
    if os.path.exists(path):
        shutil.rmtree(path)
        return True
    return False


def has_unfinished_work(template_name=None):
    """Check whether the current template still has pending or unsent work."""
    reasons = []

    if data_store.load_generation_state():
        reasons.append("generation is in progress or paused")

    pending = data_store.load_drafts(status="pending")
    if pending:
        reasons.append(f"{len(pending)} pending draft(s) waiting for review")

    approved = data_store.load_drafts(status="approved")
    if approved:
        sent_ids = data_store.get_sent_draft_ids()
        unsent = [d for d in approved if d.get("draft_id") not in sent_ids]
        if unsent:
            reasons.append(f"{len(unsent)} approved email(s) not yet sent")

    if reasons:
        return True, "; ".join(reasons)
    return False, ""


def activate_template(template_name, candidate_path, source_template_path=None, force=False):
    """Import a candidate file into an active template directory.

    Steps:
      1. Archive current template.
      2. Extract candidate to Markdown and detect source language.
      3. Call the LLM once to obtain a structured bilingual template.
      4. Write template.html (source language), template_<other>.html, and config.yaml.
      5. Mark template as unconfirmed (user must confirm before use).

    Note:
      source_template_path is accepted for CLI compatibility but is no longer used to
      merge HTML content; the new template is generated entirely from the LLM output.
    """
    if not force:
        has_work, reason = has_unfinished_work(template_name)
        if has_work:
            raise RuntimeError(
                f"Cannot import new template: current template has unfinished work ({reason}). "
                "Pass force=True or clear drafts/generation state first."
            )

    _ensure_dirs()
    _cleanup_template_images(template_name)
    archive_path = archive_current_template(template_name)

    markdown = extract_to_markdown(candidate_path, template_name=template_name)
    source_lang = detect_language(markdown, os.path.basename(candidate_path))

    # The source_template_path argument is kept for CLI compatibility but ignored;
    # the new template is produced entirely by the LLM structured output.
    _ = source_template_path

    structured = structure_template_with_llm(markdown, os.path.basename(candidate_path))
    subject_template = structured.get("subject_template", "").strip()
    if not subject_template:
        fallback = _fallback_subject(markdown)
        print(f"⚠️ 警告：模板导入返回的 subject_template 为空，已自动生成主题：{fallback}")
        structured["subject_template"] = fallback

    # Bilingual subjects: backfill missing variants from the legacy single
    # field so render-time can pick the variant matching the email language.
    subject_template = structured.get("subject_template", "").strip()
    for lang in ("cn", "en"):
        key = f"subject_template_{lang}"
        if not str(structured.get(key, "")).strip():
            structured[key] = subject_template if source_lang == lang else ""

    # Guard: English subject must stay language-pure. Drop an offending variant
    # so the render-side fallback chain produces a clean English subject.
    if re.search(r"[一-鿿]", str(structured.get("subject_template_en", ""))):
        print(
            "⚠️ 警告：英文主题包含汉字，已丢弃该变体，"
            "发送时将回退为纯英文标题。"
        )
        structured["subject_template_en"] = ""

    written = write_structured_template(template_name, structured, source_lang)

    # Mark as unconfirmed so the user has to review before generation,
    # and record the import date for the status bar.
    settings = data_store.load_settings()
    settings["template_confirmed"] = False
    settings["template_confirmed_at"] = None
    settings.setdefault("template_imported_at", {})[template_name] = datetime.now().strftime("%Y-%m-%d")
    data_store.save_settings(settings)

    return {
        "template_name": template_name,
        "source_language": source_lang,
        "archive_path": archive_path,
        "main_path": written["main_path"],
        "other_path": written["other_path"],
        "config_path": written["config_path"],
    }


# ------------------------------------------------------------------------------
# Preview
# ------------------------------------------------------------------------------

def _read_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


LANGUAGE_LABELS = {"cn": "中文版", "en": "英文版"}


def resolve_template_preview_files(template_name):
    """解析模板的双语预览文件，按「中文在前、英文在后」排序。

    语言解析与 template_engine.render() 保持一致：优先 template_<lang>.html，
    缺失或为空时回退到 template.html，并打印明确的中文提示。

    Returns:
        list of dict: {"language", "label", "path", "fallback"}。
    """
    dir_path = os.path.join(config.TEMPLATES_DIR, template_name)
    default_path = os.path.join(dir_path, "template.html")

    def _has_content(path):
        if not os.path.exists(path):
            return False
        try:
            with open(path, "r", encoding="utf-8") as f:
                return bool(f.read().strip())
        except OSError:
            return False

    variants = []
    seen_paths = set()
    for lang in ("cn", "en"):
        label = LANGUAGE_LABELS[lang]
        variant_path = os.path.join(dir_path, f"template_{lang}.html")
        if _has_content(variant_path):
            path, fallback = variant_path, False
        elif _has_content(default_path):
            path, fallback = default_path, True
            if os.path.exists(variant_path):
                print(
                    f"⚠️ 警告：模板 '{template_name}' 的 template_{lang}.html 内容为空，"
                    f"{label}预览将回退到 template.html。"
                )
            else:
                print(
                    f"ℹ️ 模板 '{template_name}' 缺少 template_{lang}.html，"
                    f"{label}预览将显示默认 template.html。"
                )
        else:
            print(
                f"⚠️ 警告：模板 '{template_name}' 没有可用的{label}文件，跳过该语言预览。"
            )
            continue
        if path in seen_paths:
            continue
        seen_paths.add(path)
        variants.append(
            {"language": lang, "label": label, "path": path, "fallback": fallback}
        )
    return variants


def build_preview_html(template_name, language=None):
    """Build a self-contained preview HTML for the active template.

    Args:
        template_name: template directory name under templates/email.
        language: optional "cn"/"en" to preview a specific language variant
            (same resolution rules as template_engine.render()).
    """
    template_path = template_engine.get_template_path(template_name, language)
    if not os.path.exists(template_path):
        html = f"<p><em>Template HTML not found for '{template_name}'.</em></p>"
    else:
        with open(template_path, "r", encoding="utf-8") as f:
            html = f.read()

    # Highlight variables and image/file placeholders
    html = re.sub(
        r"\{\{([^{}:]+)\}\}",
        r'<span style="background:#fff3cd;padding:0 4px;border-radius:3px;">{{\1}}</span>',
        html,
    )
    html = re.sub(
        r"\{\{IMAGE:([^}]+)\}\}",
        r'<div style="background:#f8d7da;padding:8px;margin:8px 0;text-align:center;border-radius:4px;">'
        r'📷 Image placeholder: \1</div>',
        html,
    )
    html = re.sub(
        r"\{\{FILE:([^}]+)\}\}",
        r'<div style="background:#d1ecf1;padding:8px;margin:8px 0;text-align:center;border-radius:4px;">'
        r'📎 File placeholder: \1</div>',
        html,
    )

    language_note = ""
    if language:
        language_note = f" — {LANGUAGE_LABELS.get(language, language)}"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Template Preview: {template_name}{language_note}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; max-width: 720px; margin: 40px auto; padding: 20px; line-height: 1.6; }}
        .banner {{ background: #d1ecf1; padding: 12px 16px; border-radius: 6px; margin-bottom: 24px; }}
    </style>
</head>
<body>
    <div class="banner"><strong>Preview mode{language_note}</strong> — variables are highlighted. Confirm in the terminal to activate this template.</div>
    {html}
</body>
</html>
"""


# ------------------------------------------------------------------------------
# Convenience: confirm active template
# ------------------------------------------------------------------------------

def confirm_active_template():
    """Mark the active template as confirmed in settings."""
    settings = data_store.load_settings()
    settings["template_confirmed"] = True
    settings["template_confirmed_at"] = datetime.now().isoformat()
    data_store.save_settings(settings)
    return settings
