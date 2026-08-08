"""阶段 G：体验优化与图片提取加固（G.1 终端 UI；G.2 双语分离预览；G.3 docx 图片提取与清理）。"""

import csv
import io
import json
import os
import struct
import sys
import types
import zlib

import pytest

from email_agent import config, data_store, preview, template_importer


@pytest.mark.parametrize("headless,expected_hint", [
    (True, "playwright install chromium"),
    (False, "未检测到本地浏览器环境"),
])
def test_playwright_failure_prints_friendly_chinese(monkeypatch, capsys, headless, expected_hint):
    """G.1：Playwright 启动失败时不得打印大段英文异常栈，只输出简短中文提示。"""
    fake_pkg = types.ModuleType("playwright")
    fake_sync = types.ModuleType("playwright.sync_api")

    def _boom(*args, **kwargs):
        raise RuntimeError(
            "Executable doesn't exist at /root/.cache/ms-playwright/chromium-1148/chrome-linux/headless_shell\n"
            "Looks like Playwright Test or Playwright is not installed.\n"
            "Please run the following command to download the browser:\n"
            "  npx playwright install chromium"
        )

    fake_sync.sync_playwright = _boom
    monkeypatch.setitem(sys.modules, "playwright", fake_pkg)
    monkeypatch.setitem(sys.modules, "playwright.sync_api", fake_sync)

    result = preview._open_with_playwright("/tmp/nonexistent_preview.html", headless=headless)
    out = capsys.readouterr().out

    assert result is False
    # 英文异常细节不得泄漏到终端
    assert "Executable doesn't exist" not in out
    assert "Looks like Playwright" not in out
    # 必须包含简短中文提示
    assert expected_hint in out


# ------------------------------------------------------------------------------
# G.2 双语模板分离预览
# ------------------------------------------------------------------------------

def _make_template(name, files):
    """Create templates/email/<name>/ with the given {filename: content} files."""
    tdir = os.path.join(config.TEMPLATES_DIR, name)
    os.makedirs(tdir, exist_ok=True)
    for fname, content in files.items():
        with open(os.path.join(tdir, fname), "w", encoding="utf-8") as f:
            f.write(content)
    return tdir


def test_resolve_preview_files_standard_layout(isolated_env):
    """标准布局（template.html 为英文源 + template_cn.html）：中英各一，中文在前。"""
    tdir = _make_template("t1", {
        "template.html": "<p>English source</p>",
        "template_cn.html": "<p>中文内容</p>",
    })
    variants = template_importer.resolve_template_preview_files("t1")
    assert [v["language"] for v in variants] == ["cn", "en"]
    assert variants[0]["path"] == os.path.join(tdir, "template_cn.html")
    assert variants[0]["fallback"] is False
    # 英文版回退到 template.html（它就是英文源）
    assert variants[1]["path"] == os.path.join(tdir, "template.html")
    assert variants[1]["fallback"] is True


def test_resolve_preview_files_missing_variant_prints_hint(isolated_env, capsys):
    """缺少语言变体文件时给出明确中文提示（复现素材：开发信测试 缺 template_en.html）。"""
    _make_template("t2", {
        "template.html": "<p>English source</p>",
        "template_cn.html": "<p>中文内容</p>",
    })
    variants = template_importer.resolve_template_preview_files("t2")
    out = capsys.readouterr().out
    assert len(variants) == 2
    assert "缺少 template_en.html" in out


def test_resolve_preview_files_empty_variant_falls_back(isolated_env, capsys):
    """语言变体文件存在但内容为空（LLM 生成缺失）→ 警告并回退 template.html。"""
    tdir = _make_template("t3", {
        "template.html": "<p>English source</p>",
        "template_cn.html": "   \n",
    })
    variants = template_importer.resolve_template_preview_files("t3")
    out = capsys.readouterr().out
    assert "内容为空" in out
    # 中英都解析到同一个 template.html，只预览一次
    assert len(variants) == 1
    assert variants[0]["path"] == os.path.join(tdir, "template.html")


def test_resolve_preview_files_no_files(isolated_env, capsys):
    """完全没有模板文件 → 返回空列表并打印警告。"""
    os.makedirs(os.path.join(config.TEMPLATES_DIR, "t4"), exist_ok=True)
    variants = template_importer.resolve_template_preview_files("t4")
    out = capsys.readouterr().out
    assert variants == []
    assert "跳过该语言预览" in out


def test_build_preview_html_language_label(isolated_env):
    """build_preview_html 支持 language 参数并在页面中标注语言。"""
    _make_template("t5", {
        "template.html": "<p>English</p>",
        "template_cn.html": "<p>中文正文</p>",
    })
    html = template_importer.build_preview_html("t5", language="cn")
    assert "中文版" in html
    assert "中文正文" in html

    html_default = template_importer.build_preview_html("t5")
    assert "English" in html_default


def test_open_template_preview_opens_both_languages(isolated_env, monkeypatch, capsys):
    """open_template_preview 必须依次打开中、英两个预览窗口。"""
    _make_template("t6", {
        "template.html": "<p>English source</p>",
        "template_cn.html": "<p>中文内容</p>",
    })
    opened = []

    def fake_open_html(html, suffix=".html"):
        opened.append(suffix)
        return f"/tmp/preview{suffix}"

    monkeypatch.setattr(preview, "_open_html", fake_open_html)
    paths = preview.open_template_preview("t6")
    out = capsys.readouterr().out

    assert len(paths) == 2
    assert opened == ["_template_cn.html", "_template_en.html"]
    assert "正在打开中文版模板预览（1/2）" in out
    assert "正在打开英文版模板预览（2/2）" in out


def test_open_template_preview_no_files(isolated_env, monkeypatch, capsys):
    """没有可用模板文件时返回空列表并给出提示，而不是抛异常。"""
    opened = []
    monkeypatch.setattr(preview, "_open_html", lambda html, suffix=".html": opened.append(suffix))
    paths = preview.open_template_preview("not_exist")
    out = capsys.readouterr().out
    assert paths == []
    assert opened == []
    assert "无法打开预览" in out


# ------------------------------------------------------------------------------
# G.2 language 生成失败容错
# ------------------------------------------------------------------------------

def _patch_llm(monkeypatch, raw_result):
    """Patch the LLM layer so structure_template_with_llm runs offline."""
    monkeypatch.setattr(
        "email_agent.config.get_active_model",
        lambda: {"name": "test", "model": "test", "api_key": "x", "base_url": ""},
    )
    monkeypatch.setattr(template_importer, "_read_file", lambda path: "SYSTEM PROMPT")
    monkeypatch.setattr(
        "email_agent.llm_client.complete_json",
        lambda system, user, schema, temperature=None: raw_result,
    )


def test_structure_llm_invalid_json_raises_chinese_message(isolated_env, monkeypatch):
    """LLM 返回非 JSON 时抛出可理解的中文错误，而不是 'Expecting value...'。"""
    _patch_llm(monkeypatch, {"content": "not a json"})
    with pytest.raises(RuntimeError, match="无法解析为 JSON"):
        template_importer.structure_template_with_llm("# Hello", filename="t.md")


def test_structure_llm_missing_content_field(isolated_env, monkeypatch):
    _patch_llm(monkeypatch, {})
    with pytest.raises(RuntimeError, match="缺少 content 字段"):
        template_importer.structure_template_with_llm("# Hello", filename="t.md")


def test_structure_llm_normalizes_alias_field_names(isolated_env, monkeypatch, capsys):
    """LLM 字段命名不符（chinese_html）时归一化为 cn_html。"""
    _patch_llm(monkeypatch, {"content": json.dumps({
        "subject_template": "S",
        "chinese_html": "<p>中文</p>",
        "en_html": "<p>English</p>",
        "variables": [],
        "images": [],
        "files": [],
    })})
    result = template_importer.structure_template_with_llm("# Hello", filename="t.md")
    out = capsys.readouterr().out
    assert result["cn_html"] == "<p>中文</p>"
    assert "已归一化为 'cn_html'" in out


def test_structure_llm_warns_on_missing_language_html(isolated_env, monkeypatch, capsys):
    """缺少某语言 HTML 时打印明确警告，供写盘阶段容错处理。"""
    _patch_llm(monkeypatch, {"content": json.dumps({
        "subject_template": "S",
        "cn_html": "",
        "en_html": "<p>English</p>",
        "variables": [],
        "images": [],
        "files": [],
    })})
    template_importer.structure_template_with_llm("# Hello", filename="t.md")
    out = capsys.readouterr().out
    assert "缺少中文模板内容" in out
    assert "'cn_html' 为空" in out


def test_write_structured_template_empty_source_raises(isolated_env):
    """源语言 HTML 缺失 → 明确报错，禁止写出空 template.html。"""
    structured = {"subject_template": "S", "en_html": "", "cn_html": ""}
    with pytest.raises(ValueError, match="模板导入失败"):
        template_importer.write_structured_template("bad", structured, "en")


def test_write_structured_template_empty_target_skips_file(isolated_env, capsys):
    """目标语言 HTML 缺失 → 不写空文件，other_path 为 None 并打印警告。"""
    structured = {"subject_template": "S", "en_html": "<p>English</p>", "cn_html": ""}
    result = template_importer.write_structured_template("t7", structured, "en")
    out = capsys.readouterr().out

    assert os.path.exists(result["main_path"])
    assert result["other_path"] is None
    assert not os.path.exists(
        os.path.join(config.TEMPLATES_DIR, "t7", "template_cn.html")
    )
    assert "cn 版模板未写入" in out

    # config.yaml 持久化源语言，供后续预览/诊断使用
    import yaml
    with open(result["config_path"], "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f)
    assert cfg["source_language"] == "en"


# ------------------------------------------------------------------------------
# G.3 docx 图片提取加固
# ------------------------------------------------------------------------------

def _png_bytes():
    """Generate a minimal valid 1x1 PNG (python-docx validates structure)."""
    def chunk(typ, data):
        return (
            struct.pack(">I", len(data)) + typ + data
            + struct.pack(">I", zlib.crc32(typ + data) & 0xFFFFFFFF)
        )
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00")
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")
    )


def test_docx_inline_image_extracted_with_placeholder(isolated_env):
    """G.3：inline 图片被提取到 assets/images/，占位符保留在原文位置。"""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Hello before")
    doc.add_paragraph().add_run().add_picture(io.BytesIO(_png_bytes()))
    doc.add_paragraph("World after")
    src = os.path.join(config.TEMPLATE_IMPORT_DIR, "src.docx")
    doc.save(src)

    markdown = template_importer._docx_to_markdown(src, template_name="t")

    assert "{{IMAGE:t_img_01}}" in markdown
    assert (
        markdown.index("Hello before")
        < markdown.index("{{IMAGE:t_img_01}}")
        < markdown.index("World after")
    )
    assert any(f.startswith("t_img_01") for f in os.listdir(config.IMAGES_DIR))


def test_docx_image_wrapped_in_alternate_content(isolated_env):
    """G.3：Word 2010+ 的 mc:AlternateContent 包裹也能提取，且 Fallback 不重复提取。"""
    import docx
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn

    doc = docx.Document()
    run = doc.add_paragraph().add_run()
    run.add_picture(io.BytesIO(_png_bytes()))
    drawing = run._element.findall(qn("w:drawing"))[0]
    rid = drawing.find(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    ).get(qn("r:embed"))

    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    v_ns = "urn:schemas-microsoft-com:vml"
    r_ns = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    )
    ac = parse_xml(
        f'<mc:AlternateContent xmlns:mc="{mc_ns}" xmlns:w="{w_ns}"'
        f' xmlns:v="{v_ns}" xmlns:r="{r_ns}">'
        '<mc:Choice Requires="w14"/>'
        f'<mc:Fallback><w:pict><v:imagedata r:id="{rid}"/>'
        "</w:pict></mc:Fallback></mc:AlternateContent>"
    )
    run._element.remove(drawing)
    ac.find(
        "{http://schemas.openxmlformats.org/markup-compatibility/2006}Choice"
    ).append(drawing)
    run._element.append(ac)

    src = os.path.join(config.TEMPLATE_IMPORT_DIR, "ac.docx")
    doc.save(src)

    markdown = template_importer._docx_to_markdown(src, template_name="t")

    # 恰好提取一次：Fallback 中的 VML 副本不得造成重复占位符
    assert markdown.count("{{IMAGE:") == 1
    assert "{{IMAGE:t_img_01}}" in markdown
    assert any(f.startswith("t_img_01") for f in os.listdir(config.IMAGES_DIR))


def test_docx_image_in_table_recovered_by_sweep(isolated_env, capsys):
    """G.3：表格内图片不在 doc.paragraphs 中，兜底扫描必须提取并给出警告。"""
    import docx

    doc = docx.Document()
    doc.add_paragraph("正文段落")
    table = doc.add_table(rows=1, cols=1)
    cell_para = table.cell(0, 0).paragraphs[0]
    cell_para.add_run().add_picture(io.BytesIO(_png_bytes()))
    src = os.path.join(config.TEMPLATE_IMPORT_DIR, "tbl.docx")
    doc.save(src)

    markdown = template_importer._docx_to_markdown(src, template_name="t")
    out = capsys.readouterr().out

    assert "{{IMAGE:t_img_01}}" in markdown
    assert "未在正文段落中定位到" in out
    assert any(f.startswith("t_img_01") for f in os.listdir(config.IMAGES_DIR))


REPRO_DOCX = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "templates", "import", "开发信带图版.docx",
)


@pytest.mark.skipif(
    not os.path.exists(REPRO_DOCX), reason="缺少人工测试素材 开发信带图版.docx"
)
def test_docx_real_repro_file_extracts_images(isolated_env):
    """G.3 验收：人工黑盒测试的含图 docx 必须提取出图片，占位符与落盘文件一一对应。"""
    markdown = template_importer._docx_to_markdown(
        REPRO_DOCX, template_name="repro"
    )

    assert "{{IMAGE:" in markdown
    extracted = [
        f for f in os.listdir(config.IMAGES_DIR) if f.startswith("repro_img_")
    ]
    assert len(extracted) >= 1
    assert markdown.count("{{IMAGE:") == len(extracted)


# ------------------------------------------------------------------------------
# G.3 图片清理保护（待审核/待发送草稿引用图片禁止删除）
# ------------------------------------------------------------------------------

def _seed_image(fname):
    path = os.path.join(config.IMAGES_DIR, fname)
    with open(path, "wb") as f:
        f.write(b"x")
    return path


def _seed_draft(draft_id, template, status, image_fname):
    drafts = data_store.load_drafts()
    drafts.append({
        "draft_id": draft_id,
        "template": template,
        "review_status": status,
        "images": [{
            "cid": os.path.splitext(image_fname)[0],
            "path": os.path.join(config.IMAGES_DIR, image_fname),
        }],
    })
    data_store.save_drafts(drafts)


def _mark_sent(draft_id):
    with open(config.EMAIL_LOGS_FILE, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=config.EMAIL_LOG_HEADERS)
        writer.writeheader()
        row = {h: "" for h in config.EMAIL_LOG_HEADERS}
        row["email_id"] = draft_id
        row["status"] = "success"
        writer.writerow(row)


def test_cleanup_protects_images_of_pending_drafts(isolated_env, capsys):
    """G.3：待审核草稿引用的图片禁止删除，未被引用的正常清理。"""
    _seed_image("t_img_01.jpg")
    _seed_image("t_img_02.jpg")
    _seed_draft("d1", "t", "pending", "t_img_01.jpg")

    template_importer._cleanup_template_images("t")
    out = capsys.readouterr().out

    assert os.path.exists(os.path.join(config.IMAGES_DIR, "t_img_01.jpg"))
    assert not os.path.exists(os.path.join(config.IMAGES_DIR, "t_img_02.jpg"))
    assert "已保留" in out


def test_cleanup_protects_images_of_approved_unsent_drafts(isolated_env):
    """G.3：已批准但未发送（待发送）草稿引用的图片禁止删除。"""
    _seed_image("t_img_01.jpg")
    _seed_draft("d1", "t", "approved", "t_img_01.jpg")

    template_importer._cleanup_template_images("t")

    assert os.path.exists(os.path.join(config.IMAGES_DIR, "t_img_01.jpg"))


def test_cleanup_removes_images_after_draft_sent(isolated_env):
    """G.3：草稿已成功发送后，对应图片可正常清理。"""
    _seed_image("t_img_01.jpg")
    _seed_draft("d1", "t", "approved", "t_img_01.jpg")
    _mark_sent("d1")

    template_importer._cleanup_template_images("t")

    assert not os.path.exists(os.path.join(config.IMAGES_DIR, "t_img_01.jpg"))


def test_cleanup_ignores_drafts_of_other_templates(isolated_env):
    """G.3：其他模板的待审核草稿不影响当前模板的图片清理。"""
    _seed_image("t_img_01.jpg")
    _seed_draft("d1", "other_template", "pending", "t_img_01.jpg")

    template_importer._cleanup_template_images("t")

    assert not os.path.exists(os.path.join(config.IMAGES_DIR, "t_img_01.jpg"))
