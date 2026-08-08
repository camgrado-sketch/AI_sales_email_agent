"""阶段 C：图片/附件自动提取、CID 内联嵌入、发送前缺失检查。"""

import base64
import os

import pytest

from email_agent import config, data_store, sender, template_importer


def _write_png(path, width=1, height=1):
    """Write a minimal valid 1x1 PNG to path."""
    # Base64 of a 1x1 red PNG
    data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVQIW2P4//8/AAX+Av4N70a4AAAAAElFTkSuQmCC"
    )
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "wb") as f:
        f.write(data)


# ------------------------------------------------------------------------------
# C.1 DOCX image extraction
# ------------------------------------------------------------------------------

def test_docx_to_markdown_extracts_image_and_inserts_placeholder(isolated_env):
    pytest.importorskip("docx")
    import docx

    png_path = os.path.join(config.IMAGES_DIR, "source.png")
    _write_png(png_path)

    docx_path = os.path.join(config.TEMPLATE_IMPORT_DIR, "with_image.docx")
    doc = docx.Document()
    doc.add_paragraph("Hello world")
    doc.add_picture(png_path)
    doc.add_paragraph("Best regards")
    doc.save(docx_path)

    markdown = template_importer._docx_to_markdown(docx_path, template_name="test_tmpl")

    assert "Hello world" in markdown
    assert "{{IMAGE:test_tmpl_img_01}}" in markdown
    assert os.path.exists(os.path.join(config.IMAGES_DIR, "test_tmpl_img_01.png"))


def test_docx_import_cleans_stale_template_images(isolated_env):
    pytest.importorskip("docx")
    import docx

    png_path = os.path.join(config.IMAGES_DIR, "source.png")
    _write_png(png_path)

    stale_image = os.path.join(config.IMAGES_DIR, "test_tmpl_img_01.png")
    _write_png(stale_image)

    docx_path = os.path.join(config.TEMPLATE_IMPORT_DIR, "fresh.docx")
    doc = docx.Document()
    doc.add_paragraph("Intro")
    doc.add_picture(png_path)
    doc.save(docx_path)

    template_importer._cleanup_template_images("test_tmpl")
    assert not os.path.exists(stale_image)

    markdown = template_importer._docx_to_markdown(docx_path, template_name="fresh_tmpl")

    assert os.path.exists(os.path.join(config.IMAGES_DIR, "fresh_tmpl_img_01.png"))
    assert "{{IMAGE:fresh_tmpl_img_01}}" in markdown


# ------------------------------------------------------------------------------
# C.2 CID inline attachment
# ------------------------------------------------------------------------------

def test_create_email_message_uses_cid_inline_images():
    png_path = os.path.join(config.IMAGES_DIR, "cid_test.png")
    _write_png(png_path)

    draft = {
        "email": "to@example.com",
        "subject": "Test",
        "html_body": "<html><body><img src='cid:hero'></body></html>",
        "text_body": "text body",
        "images": [{"cid": "hero", "path": png_path}],
    }
    msg, _ = sender.create_email_message(draft)

    image_parts = [p for p in msg.walk() if p.get_content_type().startswith("image/")]
    assert len(image_parts) == 1
    part = image_parts[0]
    assert part["Content-ID"] == "<hero>"
    assert part.get_content_disposition() == "inline"


# ------------------------------------------------------------------------------
# C.3 Pre-send missing-image check
# ------------------------------------------------------------------------------

def test_check_draft_images_reports_missing_cids():
    png_path = os.path.join(config.IMAGES_DIR, "exists.png")
    _write_png(png_path)

    drafts = [
        {
            "images": [
                {"cid": "missing", "path": "/nonexistent.png"},
                {"cid": "ok", "path": png_path},
            ]
        }
    ]
    missing = sender.check_draft_images(drafts)
    assert missing == ["missing"]


def test_process_queue_aborts_when_user_declines_missing_images(make_template, write_settings, monkeypatch, capsys):
    pytest.importorskip("docx")
    make_template("img_tmpl")
    write_settings({"selected_template": "img_tmpl", "template_confirmed": True})

    draft = {
        "draft_id": "d1",
        "customer_id": "c1",
        "email": "to@example.com",
        "subject": "Test",
        "html_body": "<p>body</p>",
        "text_body": "body",
        "images": [{"cid": "missing", "path": "/nonexistent.png"}],
        "review_status": "approved",
        "template": "img_tmpl",
    }
    data_store.save_drafts([draft])

    monkeypatch.setattr("builtins.input", lambda _: "n")
    sent = []
    monkeypatch.setattr(sender, "send_email", lambda d: sent.append(d) or True)

    sender.process_queue()

    assert len(sent) == 0
    captured = capsys.readouterr()
    assert "missing" in captured.out


def test_process_queue_sends_when_user_confirms_missing_images(make_template, write_settings, monkeypatch):
    pytest.importorskip("docx")
    make_template("img_tmpl")
    write_settings({"selected_template": "img_tmpl", "template_confirmed": True})

    draft = {
        "draft_id": "d1",
        "customer_id": "c1",
        "email": "to@example.com",
        "subject": "Test",
        "html_body": "<p>body</p>",
        "text_body": "body",
        "images": [{"cid": "missing", "path": "/nonexistent.png"}],
        "review_status": "approved",
        "template": "img_tmpl",
    }
    data_store.save_drafts([draft])

    monkeypatch.setattr("builtins.input", lambda _: "y")
    sent = []
    monkeypatch.setattr(sender, "send_email", lambda d: sent.append(d) or True)

    sender.process_queue()

    assert len(sent) == 1
    assert sent[0]["draft_id"] == "d1"
